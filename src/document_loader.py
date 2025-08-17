"""
労務相談AIエージェント - 文書読み込みモジュール

ローカル文書の読み込みと前処理を担当

機能:
- PDF, DOCX, Markdown ファイルの読み込み
- テキスト抽出とクリーニング
- メタデータ付与（ファイル名、更新日時等）
- LangChainのDocumentオブジェクトへの変換
"""

import os
import re
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

# 文書処理ライブラリ
import PyPDF2
import markdown
from docx import Document as DocxDocument

# LangChain
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

# プロジェクト設定
from config.settings import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentLoader:
    """文書読み込みクラス"""
    
    def __init__(self, chunk_size: Optional[int] = None, chunk_overlap: Optional[int] = None):
        """
        初期化
        
        Args:
            chunk_size: テキスト分割のチャンクサイズ
            chunk_overlap: チャンクの重複サイズ
        """
        self.chunk_size = chunk_size or settings.chunk_size
        self.chunk_overlap = chunk_overlap or settings.chunk_overlap
        
        # テキスト分割器の初期化
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            length_function=len,
        )
        
        # サポートされるファイル拡張子
        self.supported_extensions = {
            '.pdf': self._load_pdf,
            '.md': self._load_markdown,
            '.markdown': self._load_markdown,
            '.txt': self._load_text,
            '.docx': self._load_docx
        }
        
        logger.info(f"DocumentLoader初期化完了 - チャンクサイズ: {self.chunk_size}, 重複: {self.chunk_overlap}")
    
    def load_documents(self, folder_path: str) -> List[Document]:
        """
        指定フォルダから全ての対応文書を読み込み
        
        Args:
            folder_path: 文書フォルダのパス
            
        Returns:
            List[Document]: 読み込んだ文書のリスト
        """
        folder = Path(folder_path)
        if not folder.exists():
            logger.warning(f"フォルダが存在しません: {folder_path}")
            return []
        
        documents = []
        
        # フォルダ内の全ファイルを走査
        for file_path in folder.rglob("*"):
            if file_path.is_file() and file_path.suffix.lower() in self.supported_extensions:
                try:
                    doc = self.load_single_document(str(file_path))
                    if doc:
                        documents.extend(doc)
                        logger.info(f"読み込み完了: {file_path.name} ({len(doc)}チャンク)")
                except Exception as e:
                    logger.error(f"ファイル読み込みエラー {file_path}: {e}")
        
        logger.info(f"文書読み込み完了: {len(documents)}チャンク ({len(set(doc.metadata.get('source') for doc in documents))}ファイル)")
        return documents
    
    def load_single_document(self, file_path: str) -> List[Document]:
        """
        単一文書の読み込み
        
        Args:
            file_path: ファイルパス
            
        Returns:
            List[Document]: チャンクに分割された文書のリスト
        """
        path = Path(file_path)
        if not path.exists():
            logger.warning(f"ファイルが存在しません: {file_path}")
            return []
        
        extension = path.suffix.lower()
        if extension not in self.supported_extensions:
            logger.warning(f"サポートされていないファイル形式: {extension}")
            return []
        
        # ファイルローダーの実行
        loader_func = self.supported_extensions[extension]
        document = loader_func(str(path))
        
        if not document:
            return []
        
        # テキスト分割
        chunks = self.text_splitter.split_documents([document])
        
        # チャンクにメタデータを追加
        for i, chunk in enumerate(chunks):
            chunk.metadata.update({
                'chunk_index': i,
                'total_chunks': len(chunks),
                'chunk_size': len(chunk.page_content)
            })
        
        return chunks
    
    def _load_pdf(self, file_path: str) -> Optional[Document]:
        """PDFファイルの読み込み"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                
                for page_num, page in enumerate(pdf_reader.pages):
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"--- ページ {page_num + 1} ---\n{page_text}")
                
                if not text_content:
                    logger.warning(f"PDFからテキストが抽出できませんでした: {file_path}")
                    return None
                
                content = "\n\n".join(text_content)
                cleaned_content = self._clean_text(content)
                
                return Document(
                    page_content=cleaned_content,
                    metadata=self._get_file_metadata(file_path, 'pdf', len(pdf_reader.pages))
                )
                
        except Exception as e:
            logger.error(f"PDF読み込みエラー {file_path}: {e}")
            return None
    
    def _load_markdown(self, file_path: str) -> Optional[Document]:
        """Markdownファイルの読み込み"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            
            # MarkdownをHTMLに変換してからテキスト抽出
            html = markdown.markdown(md_content)
            # 簡単なHTMLタグ除去
            text_content = re.sub(r'<[^>]+>', '', html)
            cleaned_content = self._clean_text(text_content)
            
            return Document(
                page_content=cleaned_content,
                metadata=self._get_file_metadata(file_path, 'markdown')
            )
            
        except Exception as e:
            logger.error(f"Markdown読み込みエラー {file_path}: {e}")
            return None
    
    def _load_text(self, file_path: str) -> Optional[Document]:
        """テキストファイルの読み込み"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            cleaned_content = self._clean_text(content)
            
            return Document(
                page_content=cleaned_content,
                metadata=self._get_file_metadata(file_path, 'text')
            )
            
        except Exception as e:
            logger.error(f"テキストファイル読み込みエラー {file_path}: {e}")
            return None
    
    def _load_docx(self, file_path: str) -> Optional[Document]:
        """DOCXファイルの読み込み"""
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            
            if not paragraphs:
                logger.warning(f"DOCXからテキストが抽出できませんでした: {file_path}")
                return None
            
            content = "\n\n".join(paragraphs)
            cleaned_content = self._clean_text(content)
            
            return Document(
                page_content=cleaned_content,
                metadata=self._get_file_metadata(file_path, 'docx', len(paragraphs))
            )
            
        except Exception as e:
            logger.error(f"DOCX読み込みエラー {file_path}: {e}")
            return None
    
    def _clean_text(self, text: str) -> str:
        """テキストのクリーニング"""
        if not text:
            return ""
        
        # 余分な空白や改行の除去
        cleaned = re.sub(r'\n\s*\n\s*\n', '\n\n', text)  # 3つ以上の連続改行を2つに
        cleaned = re.sub(r'[ \t]+', ' ', cleaned)  # 連続スペース・タブを単一スペースに
        cleaned = cleaned.strip()
        
        return cleaned
    
    def _get_file_metadata(self, file_path: str, file_type: str, page_count: Optional[int] = None) -> Dict[str, Any]:
        """ファイルメタデータの取得"""
        path = Path(file_path)
        stat = path.stat()
        
        metadata = {
            'source': str(path.absolute()),
            'filename': path.name,
            'file_type': file_type,
            'file_size': stat.st_size,
            'created_at': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_at': datetime.fromtimestamp(stat.st_mtime).isoformat(),
            'loader_version': '1.0'
        }
        
        if page_count is not None:
            metadata['page_count'] = page_count
        
        return metadata
    
    def get_supported_extensions(self) -> List[str]:
        """サポートされるファイル拡張子の取得"""
        return list(self.supported_extensions.keys())
    
    def is_supported_file(self, file_path: str) -> bool:
        """ファイルがサポートされているかチェック"""
        extension = Path(file_path).suffix.lower()
        return extension in self.supported_extensions


def process_initial_documents(documents_dir: Optional[str] = None) -> List[Document]:
    """
    初期文書処理のヘルパー関数
    
    Args:
        documents_dir: 文書ディレクトリのパス（Noneの場合は設定から取得）
        
    Returns:
        List[Document]: 処理された文書のリスト
    """
    if documents_dir is None:
        documents_dir = settings.documents_dir
    
    loader = DocumentLoader()
    documents = loader.load_documents(documents_dir)
    
    logger.info(f"初期文書処理完了: {len(documents)}チャンク")
    return documents


if __name__ == "__main__":
    # テスト実行
    logging.basicConfig(level=logging.INFO)
    
    # サンプル文書の処理
    loader = DocumentLoader()
    documents = loader.load_documents(settings.documents_dir)
    
    if documents:
        print(f"✅ {len(documents)}個のチャンクを読み込みました")
        print(f"最初のチャンク: {documents[0].page_content[:200]}...")
    else:
        print("⚠️ 読み込める文書が見つかりませんでした")